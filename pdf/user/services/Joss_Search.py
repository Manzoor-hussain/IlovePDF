import os
import openai
import docx
import docx2txt
import textwrap
import re
from .keys import api_key
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT




def get_completion(prompt, model="gpt-3.5-turbo", temperature=0): 
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=temperature, 
    )
    return response.choices[0].message["content"]

def joss_converter(path, pathoutput,save_path):

    formatted = pathoutput
    #"formats/Joss_Search.docx"
    # un_formatted=os.getcwd() + "/cvs_template/Takara_Thomas.docx"
    # un_formatted = os.getcwd() + "/cvs_template/Barry_Weston_HGV.docx"
    # un_formatted = os.getcwd() + path
    # un_formatted=os.getcwd() + "/cvs_template/Nicholas_Eager.docx"


    # extract the text from the Word document
    doc = docx.Document(path)
    formated_text = docx2txt.process(formatted)
    unformated_text = docx2txt.process(path)
    
    print("Process has started...")

    fields_labels = "Name, Notice Period, Holiday Dates, Candidate Overview, Summary, Experience, Education, Courses, Previous Assignments,Professional Qualifications, Areas of Expertise, Key Skills, Computer Skills, Languages, Interests"
    test_text=f"""

    Perform the following tasks on this text: "{unformated_text}":

    1: Extract the information according to these labels {fields_labels}.

    2: Must include all the fields and subfields i.e. fields_labels in the output. 

    3: The output must be in key value format.

    4: Do not include any fields or subfields other than the mentioned.

    5: Do not include email, phone number and personal address.

    6: The following field labels should be in tabular format: Notice Period, Holiday Dates, Candidate Overview.

    7: Format the Key Skills in bullet points.

    8: Format Experience in bullet points

    8: Must mention the proficiency of the language as well if found e.g. English (Native)

    9: Must display Job Title in new line under Company Name in Experience

    10: Summary must be in paragraph format

    11: The template should look like this:


    Name



    Notice Period: 

    Holiday Dates: 

    Candidate Overview:




    Summary



    Experience


    Education (Institution name and duration should be on the same line)

        For example (This is just to give you an example do not display the given lines below)
        Oxford University, London, March 2014 - May 2016
        Name of first degree
        Name of second degree
        Name of third degree


    Courses


    Previous Assignments


    Professional Qualifications


    Areas of Expertise


    Key Skills
        Add - before every key skill

    Computer Skills


    Languages
    Example: English (Native)

    Interests


    Experience should look like this:

        Company Name, Duration (Company name and duration should be in same line) 
        Job Title
        Responsibilities
    """
    # text=f"{test_text}"
    # result=llm(text)
    result = get_completion(test_text)
    print ("Results are: ", result)
    txt='\n'+result.replace("\n","\n\n") + '\n'
    name_pattern = r'(\nName ?:|\nNotice Period ?:|\nHoliday Dates ?:|\nCandidate Overview ?:|\nSummary ?:|\nExperience ?:|\nEducation ?:|\nCourses ?:|\nPrevious Assignments ?:|\nProfessional Qualifications ?:|\nAreas of Expertise ?:|\nKey Skills ?:|\nComputer Skills ?:|\nLanguages ?:|\nInterests ?:)'

    try:
        name = re.split(name_pattern,txt)
    except:
        name = ''

    dc = {i:'' for i in ['Name', 'Notice Period', 'Holiday Dates', 'Candidate Overview', 'Summary', 'Experience', 'Education', 'Courses', 'Previous Assignments', 'Professional Qualifications','Areas of Expertise','Key Skills','Computer Skills','Languages','Interests']}
    for ind,i in enumerate(name):
        try:
            if i.strip(' \n:') in dc and name[ind+1].strip(' :') not in dc:
                dc[i.strip(' \n:')] += '\n' + name[ind+1].strip()

        except:
            pass
    print('DICT')
    print("Dictionary is: ", dc)
    # Open the existing document
    doc = docx.Document(formatted)

    # Get the first paragraph
    for i,p in enumerate(doc.paragraphs):
        for key in dc:
            if p.text.strip(' :\n').lower() == key.lower():
    #             print(key)



                if key.lower() in ['experience']:
                    conv_text = re.sub('[\n ]*\n[\n ]*- *','\n• ',dc[key])
                    conv_list = re.split('\n',re.sub('\n? *\n\n *| *\n *','\n',conv_text))
                
                    for b in conv_list:
                        if '• ' in b:
                            doc.paragraphs[i+2].add_run(b.strip()+ '\n')
                        else:
                            doc.paragraphs[i+2].add_run(b.strip() + '\n').bold = True

                            
                elif key.lower() in ['education']:
                    for c in re.split('\n',re.sub('\n? *\n\n *| *\n *','\n',dc[key].replace(';',',').replace(':','\n• ').replace(' - ', '-').replace('- ', '• ').replace(' -','• '))):
                        if '• ' in c:
                            doc.paragraphs[i+2].add_run(c.strip()+ '\n')
                        elif c.startswith("-"):
                            doc.paragraphs[i+2].add_run(c.strip()+ '\n')
                        else:
                            doc.paragraphs[i+2].add_run(c.strip() + '\n').bold = True

    
    
                elif key.lower() in ['key skills']:
                    formatted_text = ''
                    conv_text = re.sub(';|-|,','\n', str(dc[key]))
                    conv_text = re.sub('[ \n]*\n[ \n]*','\n', conv_text)
                    groups = re.split('\n',conv_text)
                    for group in groups:
                        if len(groups) == 1:
                            formatted_text = str(dc[key])
                            break
                        wrapper = textwrap.TextWrapper(width=60, initial_indent='• ',
                                                       subsequent_indent='  ')
                        formatted_text += wrapper.fill(group) + '\n'
                    doc.paragraphs[i+2].text = formatted_text.strip()
                    
                    

                elif key.lower() in ['areas of expertise']:
                    formatted_text = ''
                    conv_text = re.sub(';|-|,','\n', str(dc[key]))
                    conv_text = re.sub('[ \n]*\n[ \n]*','\n', conv_text)
                    groups = re.split('\n',conv_text)
                    for group in groups:
                        if len(groups) == 1:
                            formatted_text = str(dc[key])
                            break
                        wrapper = textwrap.TextWrapper(width=60, initial_indent='• ',
                                                       subsequent_indent='  ')
                        formatted_text += wrapper.fill(group) + '\n'
                    doc.paragraphs[i+2].text = formatted_text.strip()
    #                 doc.paragraphs[i+2].text = dc[key].replace(';',',').replace('\n- ','• ').replace(',','\n• ')
                
                elif key.lower() in ['computer skills']:
                    formatted_text = ''
                    conv_text = re.sub(';|-|,','\n', str(dc[key]))
                    conv_text = re.sub('[ \n]*\n[ \n]*','\n', conv_text)
                    groups = re.split('\n',conv_text)
                    for group in groups:
                        if len(groups) == 1:
                            formatted_text = str(dc[key])
                            break
                        wrapper = textwrap.TextWrapper(width=60, initial_indent='• ',
                                                       subsequent_indent='  ')
                        formatted_text += wrapper.fill(group) + '\n'
                    doc.paragraphs[i+2].text = formatted_text.strip()
                    
                    

                elif key.lower() in ['languages','interests', 'professional qualifications', 'achievements', 'courses']:

                    formatted_text = ''
                    groups = re.split(',|;|-',str(dc[key]))
                    for group in groups:
    #                     print(groups)
    #                     break

                        if len(groups) == 1:
                            formatted_text = str(dc[key])
                            break
                        wrapper = textwrap.TextWrapper(width=60, initial_indent='• ',
                                                       subsequent_indent='  ')
                        formatted_text += wrapper.fill(group) + '\n'
                    doc.paragraphs[i+2].text = formatted_text.strip()
                elif key.lower() in ['Summary']:
                    doc.paragraphs[i+2].text = dc[key].replace('\n\n\n\n',',')+"\n\n"
                else:
                    doc.paragraphs[i+2].text = str(dc[key])

    #             else:
    #                 doc.paragraphs[i+2].text = str(dc[key])

    for table in doc.tables:
        for row in table.rows:
            for i,cell in enumerate(row.cells):
                for key in dc:
                    if cell.text.strip(' :\n').lower() == key.lower().replace('current ',''):
                        row.cells[i+1].text = str(dc[key])

    # Save the updated document as a new file
    doc.save(save_path)

    print("Conversion Completed...")
