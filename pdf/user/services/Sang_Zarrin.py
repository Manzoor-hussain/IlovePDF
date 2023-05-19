import os
import openai
import docx
import docx2txt
import re
import json
from .keys import api_key
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def sang_zarrin_converter(path, pathoutput,save_path):
    formatted= pathoutput
   # un_formatted=os.getcwd() + path
    # un_formatted=os.getcwd() + "/unformated/AymanAbouChakra_CV_2021.docx"
    # un_formatted=os.getcwd() + "/unformated/Ayman Abdul-Hadi.docx"
    # un_formatted=os.getcwd() + "/unformated/Amr Abdelbaki.docx"


    doc = docx.Document(path)
    formated_text = docx2txt.process(formatted)
    unformated_text = docx2txt.process(path)


    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    # llm=OpenAI(temperature=0, max_tokens=1500,openai_api_key=api_key)
    
    
    print("Process has Started...")
    fields_labels = "Name, SUMMARY, EXPERIENCE, EDUCATION, CERTIFICATIONS, TRAININGS, COMPUTER SKILLS, SKILLS, QUALIFICATIONS,   LANGUAGES, INTERESTS"

    
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Summary" : "value",
    "Experience" : [
        {"Designation" : "The specific designation or position on which he works in this company",
        "Company Name" : ["Name of company", "Location of company"],
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Company Location" : "Location of company",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ]
    "Education" : [
        {"Institute Name" : "Name Of institute and its location if available separated with comma "," ",
        "Degree Nmae": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Nmae": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        ...
        ],
    "Trainings" : ["trainings1", "trainings2", ...],
    "Computer Skills" : ["computer skill1", "computer skill2", ...],
    "Skills" : ["skill1", "skill2", ...],
    "Qualifications" : ["qualification1", "qualification2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["interest1", "interest2", ...]
    }
    """


    result = get_completion(test_text)


    dc = dict(json.loads(re.sub(',[ \n]*\]',']',re.sub(',[ \n]*\}','}',result.replace('...','')))))


    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        
        try:
            if p.text.strip(' :\n').lower() == 'summary':
                doc.paragraphs[i+2].add_run(dc['Summary'].strip()).bold = False
                doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'experience':
                for j in dc['Experience']:
                    doc.paragraphs[i+2].add_run(j['Designation'].strip() + '\n').bold = True
                    doc.paragraphs[i+2].add_run(j["Company Name"][0].strip() + ',  ' + j["Company Name"][1].strip() + '\n').bold = True
                    doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n\n').bold = True
                    doc.paragraphs[i+2].add_run('Responsibilities:' + '\n').bold = True
                    for k in j['Responsibilities']:
                        doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                    doc.paragraphs[i+2].add_run('\n\n')
        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
                    doc.paragraphs[i+2].add_run(j['Institute Name'].strip() + '\n').bold = False
                    doc.paragraphs[i+2].add_run(j['Degree Nmae'].strip() + '\n').bold = False
                    doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n\n').bold = False

        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'trainings':
                for j in dc['Trainings']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'computer skills':
                for j in dc['Computer Skills']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'skills':
                for j in dc['Skills']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'qualifications':
                for j in dc['Qualifications']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'languages':
                for j in dc['Languages']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'interests':
                for j in dc['Interests']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass
    doc.save(save_path)
    print("\n")
    print("--------------------------------------------------------------------------------------------------------------------")
    print("\n")
    print("Process has Completed...")

# path="/Users/manzoorhussain/Documents/Services/IlovePDF/pdf/media/pdf_input/Takara_Thomas.docx"
# pathoutput="/Users/manzoorhussain/Documents/Services/IlovePDF/pdf/media/pdf_output/Sang_template.docx"
# save_path="/Users/manzoorhussain/Documents/Services/IlovePDF/pdf/media/pdf_input"
# Sang_Converter(path, pathoutput,save_path)