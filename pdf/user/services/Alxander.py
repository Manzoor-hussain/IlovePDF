import os
import openai
import docx
import docx2txt
import re
import json
from .keys import api_key
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def alexander_steele_converter(path,pathoutput,save_path):
    formatted= pathoutput
    # un_formatted=os.getcwd() + "/Unformated_EdEx/Adil Thomas Daraki CV .docx"
    # un_formatted=os.getcwd() + "/Unformated_EdEx/Alice Maynard CV.docx"
   
    # un_formatted=os.getcwd() + "/Unformated_EdEx/Amrit Bassan CV.docx"
    
    doc = docx.Document(path)
    formated_text = docx2txt.process(formatted)
    unformated_text = docx2txt.process(path)

    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    # llm=OpenAI(temperature=0, max_tokens=1500,openai_api_key=api_key)
    fields_labels = "Name, PROFILE, EDUCATION, IT LITERACY, CERTIFICATES, PROJECTS, PROFESSIONAL QUALIFICATIONS, SOFTWARES, languages, Interests, TRAININGS, skills, WORK EXPERIENCE"

    
    print ("Process has Started...")
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Profile" : "value",
    "Experience/Employment History" : [
        {"Company Name" : ["Name of company", "Working Duration in Company"],
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : ["Name of company", "Working Duration in Company"],
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ]
    "Education" : [
        {"Institute Name" : ["Name Of institute","Studying duration in institute"],
        "Degree Name": "Name of degree",
        },
        {"Institute Name" : ["Name Of institute", "Studying duration in institute"],
        "Degree Name": "Name of degree", 
        },
        ...
        ],
    "Professional Training/Courses" : ["Training1", "Training2", ...],
    "Achievements" : ["achievement1", "achievement2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["interest1", "interest2", ...],
    "Computer Skills" : ["computerskill1", "computerskill2", ...],
    "Skills" : ["skill1", "skill2", ...],

    }

    Do not include Grade

    Do not include Mobile number, Emali and home address 
    """


    result = get_completion(test_text)
    print(result)
    dc = dict(json.loads(re.sub(',[ \n]*\]',']',re.sub(',[ \n]*\}','}',result.replace('...','')))))

    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        try:
            if p.text.strip(' :\n').lower() == 'profile':
                doc.paragraphs[i+2].add_run(dc['Profile'].strip()).bold = False
                doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        try:        
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
        #             doc.paragraphs[i+2].add_run(j['Institute Name']).bold = Fals
                    doc.paragraphs[i+2].add_run(j["Institute Name"][0].strip() + ' – ' + j["Institute Name"][1].strip() + '\n').font.underline = False
                    doc.paragraphs[i+2].add_run(j['Degree Name'].strip() + '\n\n').bold = True
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'professional training':
                for j in dc['Professional Training/Courses']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass


        try:
            if p.text.strip(' :\n').lower() == 'achievements':
                for j in dc['Achievements']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'computer skills':
                for j in dc['Computer Skills']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass


        try:
            if p.text.strip(' :\n').lower() == 'languages':
                for j in dc['Languages']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                for j in dc['Interests']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'trainings':
                for j in dc['Trainings']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                for j in dc['Skills']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'experience':
                for j in dc['Experience/Employment History']:
                    doc.paragraphs[i+2].add_run(j['Company Name'][0].strip() + ' – ' + j['Company Name'][1] + '\n').font.underline = False
                    doc.paragraphs[i+2].add_run(j['Designation'].strip() + '\n\n').bold = True
                    doc.paragraphs[i+2].add_run('Responsibilities:' + '\n').bold = True
                    for k in j['Responsibilities']:
                        doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                    doc.paragraphs[i+2].add_run('\n\n')
        except:
            pass

        doc.save(save_path)
    print("\n")
    print("---------------------------------------------------------------------------------------------------------------------")
    print("\n")
    print("Process has Completed...")
path = "/Unformated_Fair/Alex Betts CV Ashbys.docx"
