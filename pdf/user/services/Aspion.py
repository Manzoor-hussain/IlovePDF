import os
import openai
import docx
import docx2txt
import textwrap
import re
from .keys import api_key
import re
import textwrap


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]



def aspion_converter(path, pathoutput,save_path):
    formatted = pathoutput
    #"formats/Aspion.docx"
    # un_formatted=os.getcwd() + path

    doc = docx.Document(path)
    formated_text = docx2txt.process(formatted)
    unformated_text = docx2txt.process(path)

    openai.api_key = api_key
    
    print("Process Started...")

    fields_labels = "Summary, Career History, Education, Training, Skills, Achievements, Languages, Interests"
    test_text=f"""

    Do the following tasks on this text: "{unformated_text}":

    1: Extract the information according to these labels {fields_labels}.

    2: Must include all the fields and subfields i.e. fields_labels in the output. 

    3: The output must be in key value format.

    4: Do not include any fields or subfields other than the mentioned.

    5: Do not include email, phone number and personal address.

    6: Do not include Name, Job title, Location, Salary, Notice

    7. Must add one linespace '\n' after job duration in career history.

    8. Do not include ',' between Comapny name and company locality however put '-'.

    9. Add reponsibilities if founded, below the duration in the field of career history line by line and add %- sign before every responsibility and replace every , with % in reponsibilities. 
    For example
    %- first reponsibility 
    %- second responsibility
    %- third responsibility


    10. The template should look like this:



    Summary:

        Extract only those information which describe the resume. Write them in the form of paragraph.

    Career History:

        Write Working Designation
        Write Company name, Company Locality/Place (should be in same line and not included "," however put "-")
        Write Time Period for a particular designation (and put $ before Time Period for a particular designation)

    Education:

        In education field, write name of institue and it's locality in same line, after that just write the title of
        degree completed form this institute in next line. After compliting this, write the time period of that degree 
        in next line.
        For example
        Oxford University, London 
        Becholars in Computer Science
        Oct 2008 - Nov 2010

        If you find comma before time period or title degree that was completed from this institute than replace it
        with "$".

    Training:


    Achievements:


    Skills:


    Languages:


    Interests:


    in carreer histroy include all subfields information same as below format if available but without
    including label or heading and sequence must be same as below.
    career histroy:
        Working Designation 
        Company name, Company Locality/Place (should be in same line and not included "," however put "-")
        Time Period for a particular designation after time period double line space must be given and replace "," or "-" with "$"
        if founded before Time Period for a particular designation 


        Responsibilities
    for example
        career histroy:
            Data scientist 
            ABC company,london 
            2018-2022






    """

    result = get_completion(test_text)

    print(result)
    txt='\n'+result.replace("\n","\n\n") + '\n'
    name_pattern = r'(Summary ?:|Career History ?:|Education ?:|Training ?:|Skills ?:|Achievements ?:|Languages ?:|Interests ?:)'

    try:
        name = re.split(name_pattern,txt)
    except:
        name = ''

    dc = {i:'' for i in ['Summary', 'Career History', 'Education', 'Training', 'Skills', 'Achievements', 'Languages', 'Interests']}
    for ind,i in enumerate(name):
        try:
            if i.strip(' :') in dc and name[ind+1].strip(' :') not in dc:
                dc[i.strip(' :')] += name[ind+1].strip()
        except:
            pass
        
     
    print(dc)

    dc['Skills'] = dc['Skills'].replace("\n","")
    dc['Languages'] = dc['Languages'].replace("\n\n","\n")



    # Open the existing document
    doc = docx.Document(formatted)

    # Get the first paragraph
    for i,p in enumerate(doc.paragraphs):
        for key in dc:
            if p.text.strip(' :\n').lower() == key.lower().replace('current ',''):

                if key.lower() in ['career history']:
                    for b in re.split('\n',re.sub('\n? *\n\n *| *\n *','\n',dc[key].replace('%-', '• '))):
                        if '• ' in b:
                            doc.paragraphs[i+2].add_run(b.strip()+ '\n')
                        elif b.startswith("-"):
                            doc.paragraphs[i+2].add_run(b.strip().replace("-","• ")+ '\n')
                        else:
                            doc.paragraphs[i+2].add_run(b.strip().replace('$',"") + '\n').bold = True

                elif key.lower() in ['education']:
                    doc.paragraphs[i+2].text = re.sub('\n? *\n\n *| *\n *','\n',dc[key])
                elif key.lower() in ['training']:
                    doc.paragraphs[i+2].text = re.sub('\n? *\n\n *| *\n *','\n',dc[key].replace(',','\n').replace(';','\n'))
                elif key.lower() in ['languages','interests', 'skills', 'achievements']:
                    formatted_text = ''
                    groups = re.split(',|;|%-|\n|-',str(dc[key]))
                    for group in groups:
                        if len(groups) == 1:
                            formatted_text = str(dc[key])
                            break
                        wrapper = textwrap.TextWrapper(width=60, initial_indent='• ',
                                                       subsequent_indent='  ')
                        formatted_text += wrapper.fill(group.strip()) + '\n'
    #                     print("dd",formatted_text)
                    try:  
                        doc.paragraphs[i+2].text = formatted_text.strip()
                    except:
                        pass
                else:
                    doc.paragraphs[i+2].text = str(dc[key])

    for table in doc.tables:
        for row in table.rows:
            for i,cell in enumerate(row.cells):
                for key in dc:
                    if cell.text.strip(' :\n').lower() == key.lower().replace('current ',''):
                        row.cells[i+1].text = str(dc[key])

    # Save the updated document as a new file
    doc.save(save_path)
    print("-------------------------------------------------------------------------------------------------------------------")
    print("Process Completed...")
