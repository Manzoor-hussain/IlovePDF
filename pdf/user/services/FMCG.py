import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]


def fmcg_converter(path, pathoutput,save_path):
    
    formatted= pathoutput
    
  
    
    # extract the text from the Word document
    doc = docx.Document(path)
    formatted_text = docx2txt.process(formatted)
    unformatted_text = docx2txt.process(path)
    
    
    print("Process has started...")
    
    # Prompt
    openai.api_key = api_key

    test_text = """

    Extract data from this text:

    \"""" + unformatted_text + """\"

    in following JSON format:
    {
    "Current Employer" : "value",
    "Job title" : "value",
    "Location" : "value",
    "Salary Sought" : "value",
    "Notice Period" : "value",

    "Name" : "value",
    "Profile" : "value",
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        ...
        ],
    "Professional Qualifications" : ["Qualification1", "Qualification2", ...],
    "Skills" : ["Skill1", "Skill2", ...],
    "IT Skills" : ["IT Skill1", "IT Skill2", ...],
    "Activities" : ["Activity1", "Activity2", ...],
    "Interests" : ["interest1", "interest2", ...],
    "Languages" : ["Language1", "Language2", ...],
    "Employment Summary" : [
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ]
    }
    make it sure to keep the response in JSON format.
    """

    result = get_completion(test_text)
#     print(result)
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
 
    
    doc = docx.Document(formatted)

    for table in doc.tables:
        for row in table.rows:
            for i,cell in enumerate(row.cells):
                try:
                    if cell.text.strip(' :\n').lower() == 'current employer':
                        row.cells[i+1].text = dc['Current Employer']
                except:
                    pass
                try:
                    if cell.text.strip(' :\n').lower() == 'job title':
                        row.cells[i+1].text = dc['Job Title']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'location':
                        for j in dc['Location']:
                            row.cells[i+1].text = row.cells[i+1].text + j
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'salary sought':
                        row.cells[i+1].text = dc['Salary Sought']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'notice period':
                        row.cells[i+1].text = dc['Notice Period']
                except:
                    pass                



    for i,p in enumerate(doc.paragraphs):


        if p.text.strip(' :\n').lower() == 'name':
            try:
                name_paragraph = doc.paragraphs[i]
                name_paragraph.text = str(dc['Name'])
                name_paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                name_paragraph.runs[0].bold = True
            except:
                pass


        if p.text.strip(' :\n').lower() == 'profile':
            try:
                doc.paragraphs[i+2].text = str(dc['Profile'])
                doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    institute_name = j['Institute Name'].strip()
                    duration = j['Duration'].strip()
                    degree_name = j['Degree Name'].strip()

                    doc.paragraphs[i+2].add_run(institute_name + ' ').bold = True
                    if duration:
                        doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                    else:
                        doc.paragraphs[i+2].add_run('(' + "Not mentioned" + ')' + '\n').bold = True
                    if degree_name:
                        doc.paragraphs[i+2].add_run(degree_name + '\n\n').bold = False
                    else:
                        doc.paragraphs[i+2].add_run("Not mentioned" + '\n\n').bold = False 
            except:
                pass


        if p.text.strip(' :\n').lower() == 'professional qualifications':
            try:
                for j in dc['Professional Qualifications']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass


        if p.text.strip(' :\n').lower() == 'skills':
            try:
                for j in dc['Skills']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


        if p.text.strip(' :\n').lower() == 'it skills':
            try:
                for j in dc['IT Skills']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass        



        if p.text.strip(' :\n').lower() == 'activities':
            try:
                for j in dc['Activities']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass

        if p.text.strip(' :\n').lower() == 'interests':
            try:
                for j in dc['Interests']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass

        if p.text.strip(' :\n').lower() == 'languages':
            try:
                for j in dc['Languages']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'employment summary':
            try:
                for j in dc['Employment Summary']:
                    company_name = j['Company Name'].strip()
                    duration = j['Duration'].strip()
                    job_title = j['Job Title'].strip()

                    doc.paragraphs[i+2].add_run(company_name + ' ').bold = True
                    doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                    doc.paragraphs[i+2].add_run(job_title + '\n\n').bold = False
    #                 doc.paragraphs[i+2].add_run('Duties:' + '\n\n')
                for k in j['Responsibilities']:
                    doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n')
                    doc.paragraphs[i+2].add_run('\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


    doc.save(save_path)
    print("Conversion has completed !!")