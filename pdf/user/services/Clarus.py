import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]


def clarus_converter(path,pathoutput,save_path):
    
    
    formatted = pathoutput 
    # extract the text from the Word document
    doc = docx.Document(path)
    formatted_text = docx2txt.process(formatted)
    unformatted_text = docx2txt.process(path)
    
    
    print("Process has started...")
    
    # Prompt
    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key

    test_text = """

    Ectract data from this text:

    \"""" + unformatted_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Personal Statement" : "value",
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        ...
        ],
    "Qualification" : ["Qualification1", "Qualification2", ...],
    "Employment Summary" : [
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ],
    "Other Experience" : ["Other Experience1", "Other Experience2", ...],
    "Projects and Exhibitions" : ["Projects and Exhibitions1", "Projects and Exhibitions2", ...],
    "Voluntary Experience/Work" : ["Voluntary Experience/Work1", "Voluntary Experience/Work2", ...],
    "Skills and Hobbies" : ["Skills and Hobbies1", "Skills and Hobbies2", ...],
    "Languages" : ["Language1", "Language2", ...],
    "Leadership" : ["Leadership1", "Leadership2", ...],
    "Interests" : ["interest1", "interest2", ...]
    }

    """
    # Prompt result
    result = get_completion(test_text)
    
    dc = dict(json.loads(re.sub(',[ \n]*\]',']',re.sub(',[ \n]*\}','}',result.replace('...','')))))
    
    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):


        if p.text.strip(' :\n').lower() == 'name':
            try:
                name_paragraph = doc.paragraphs[i+2]
                name_paragraph.text = str(dc['Name'])
                name_paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
            except:
                pass


        if p.text.strip(' :\n').lower() == 'personal statement':
            try:
                doc.paragraphs[i+2].text = str(dc['Personal Statement'])
                doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    institute_name = j['Institute Name'].strip()
                    duration = j['Duration'].strip()
                    degree_name = j['Degree Name'].strip()

                    doc.paragraphs[i+2].add_run(institute_name + ' ').bold = True
                    doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                    doc.paragraphs[i+2].add_run(degree_name + '\n\n').bold = False
            except:
                pass


        if p.text.strip(' :\n').lower() == 'qualifications':
            try:
                for j in dc['Qualifications']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass


        if p.text.strip(' :\n').lower() == 'employment summary':
            try:
                for j in dc['Employment Summary']:
                    company_name = j['Company Name'].strip()
                    duration = j['Duration'].strip()
                    job_title = j['Job Title'].strip()

                    doc.paragraphs[i+2].add_run(company_name + ' ').bold = True
                    doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                    doc.paragraphs[i+2].add_run(job_title + '\n\n').bold = False
    #                 doc.paragraphs[i+2].add_run('Duties:' + '\n\n')
                for k in j['Responsibilities']:
                    doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n')
                    doc.paragraphs[i+2].add_run('\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


        if p.text.strip(' :\n').lower() == 'other experience':
            try:
                for j in dc['Other Experience']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


        if p.text.strip(' :\n').lower() == 'projects and exhibitions':
            try:
                for j in dc['Projects and Exhibitions']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass

        if p.text.strip(' :\n').lower() == 'voluntary experience/work':
            try:
                for j in dc['Voluntary Experience/Work']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass

        if p.text.strip(' :\n').lower() == 'skills and hobbies':
            try:
                for j in dc['Skills and Hobbies']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass

        if p.text.strip(' :\n').lower() == 'languages':
            try:
                for j in dc['Languages']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'leadership':
            try:
                for j in dc['Leadership']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass

        if p.text.strip(' :\n').lower() == 'Interests':
            try:
                for j in dc['interests']:
                    doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


    doc.save(save_path)
    print("Conversion has completed !!")
    
    
    