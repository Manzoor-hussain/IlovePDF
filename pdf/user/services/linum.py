import os
import re
import json
from pprint import pprint

import docx
import docx2txt
import textwrap
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import openai
from .keys import api_key
openai.api_key = api_key


def get_completion(prompt, model="gpt-3.5-turbo", temperature=0): 
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=temperature, 
    )
    return response.choices[0].message["content"]

def linum_converter(path,pathoutput,save_path):

    formatted = pathoutput
    #un_formatted = os.getcwd() + path


    # extract the text from the Word document
    formated_text = docx2txt.process(formatted)
    unformated_text = docx2txt.process(path)
    
    print("Process has started...")

    test_text = """
    Extract data from this text:

    \"""" + re.sub('\n+','\n', unformated_text) + """\"

    in following JSON format:
    {
    "Candidate Name" : "value",
    "Position Applied For" : "value",
    "Notice Period" : "value",
    "Current Location" : "value",
    "Current Salary" : "value",
    "Expected Salary" : "value",
    "Marital Status" : "value",
    "Nationality" : "value",
    "Date of Birth" : "value",
    "Summary" : "value",
    "Work Experience" : [
        {"Company Name" : "Name of company",
        "Company Location" : "Location of company",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Company Location" : "Location of company",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ]
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        ...
        ],
    "Achievements" : ["achievement 1", "achievement 2", ...]
    "Qualifications" : ["qualification 1", "qualification 2", ...],
    "Skills" : ["skill 1", "skill 2", ...],
    "Attributes" : ["attribute 1", "attribute 2", ...],
    "Languages" : ["language 1", "language 2", ...],
    "Interests" : ["interest 1", "interest 2", ...]
    }
    """

    result = get_completion(test_text)
    print ("RESULTS\n\n" )
    print(result)
    print('\n\nRS_END')
    
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    print('DICTIONARY\n\n')
    print(dc)
    print('\n\nDC_END')
    
    
    # Open the existing document
    doc = docx.Document(formatted)

    # Get the first paragraph
    for i,p in enumerate(doc.paragraphs):
        if p.text.strip(' :\n').lower() == 'summary':
            try:
                doc.paragraphs[i+2].add_run(dc['Summary']).bold = False
            except:
                pass
        if p.text.strip(' :\n').lower() == 'skills':
            try:
                for j in dc['Skills']:
                    doc.paragraphs[i+2].add_run('    • ' + j.strip() + '\n').bold = False
            except:
                pass
        if p.text.strip(' :\n').lower() == 'attributes':
            try:
                for j in dc['Attributes']:
                    doc.paragraphs[i+2].add_run('    • ' + j.strip() + '\n').bold = False
            except:
                pass
        if p.text.strip(' :\n').lower() == 'achievements':
            try:
                for j in dc['Achievements']:
                    doc.paragraphs[i+2].add_run('    • ' + j.strip() + '\n').bold = False
            except:
                pass        
        if p.text.strip(' :\n').lower() == 'languages':
            try:
                for j in dc['Languages']:
                    doc.paragraphs[i+2].add_run('    • ' + j.strip() + '\n').bold = False
            except:
                pass     
        if p.text.strip(' :\n').lower() == 'interests':
            try:
                for j in dc['Interests']:
                    doc.paragraphs[i+2].add_run('    • ' + j.strip() + '\n').bold = False
            except:
                pass        
        if p.text.strip(' :\n').lower() == 'work experience':
            try:
                for j in dc['Work Experience']:
                    try:
                        if j['Duration'].strip():
                            doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n').bold = False
                        else:
                            doc.paragraphs[i+2].add_run('Duration not mentioned\n').bold = False
                    except:
                        doc.paragraphs[i+2].add_run('Duration not mentioned\n').bold = False
                    try:
                        if j['Company Name'].strip():
                            doc.paragraphs[i+2].add_run(j['Company Name'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run('Company Name not mentioned\n').bold = True
                    except:
                        doc.paragraphs[i+2].add_run('Company Name not mentioned\n').bold = True
                    try:
                        if j['Company Location'].strip():
                            doc.paragraphs[i+2].add_run(j['Company Location'].strip() + '\n\n').bold = False
                        else:
                            doc.paragraphs[i+2].add_run('Company Location not mentioned\n\n').bold = False
                    except:
                        doc.paragraphs[i+2].add_run('Company Location not mentioned\n\n').bold = False

                    try:
                        if j['Responsibilities']:
                            doc.paragraphs[i+2].add_run('Duties:' + '\n').bold = False
                            for k in j['Responsibilities']:
                                doc.paragraphs[i+2].add_run('    • ' + k.strip() + '\n').bold = False
                    except:
                        pass
                    doc.paragraphs[i+2].add_run('\n\n').bold = False
            except:
                pass

        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    try:
                        if j['Duration'].strip():
                            doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n').bold = False
                        else:
                            doc.paragraphs[i+2].add_run('Duration not mentioned\n').bold = False
                    except:
                        doc.paragraphs[i+2].add_run('Duration not mentioned\n').bold = False
                    try:
                        if j['Degree Name'].strip():
                            doc.paragraphs[i+2].add_run(j['Degree Name'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run('Degree Name not mentioned\n').bold = True
                    except:
                        doc.paragraphs[i+2].add_run('Degree Name not mentioned\n').bold = True
                    try:
                        if j['Institute Name'].strip():
                            doc.paragraphs[i+2].add_run(j['Institute Name'].strip() + '\n\n').bold = False
                        else:
                            doc.paragraphs[i+2].add_run('Institute Name not mentioned\n\n').bold = False
                    except:
                        doc.paragraphs[i+2].add_run('Institute Name not mentioned\n\n').bold = False
            except:
                pass



    for table in doc.tables:
        for row in table.rows:
            for i,cell in enumerate(row.cells):
                try:
                    if cell.text.strip(' :\n').lower() == 'position applied for':
                        row.cells[i+1].text = dc['Position Applied For']
                except:
                    pass
                try:
                    if cell.text.strip(' :\n').lower() == 'current location':
                        row.cells[i+1].text = dc['Current Location']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'qualifications':
                        for j in dc['Qualifications']:
                            row.cells[i+1].text = row.cells[i+1].text + j + '\n'
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'current salary':
                        row.cells[i+1].text = dc['Current Salary']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'expected salary':
                        row.cells[i+1].text = dc['Expected Salary']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'availability / notice period':
                        row.cells[i+1].text = dc['Notice Period']
                except:
                    pass
                try:
                    if cell.text.strip(' :\n').lower() == 'personal details':
                        try:
                            row.cells[i+1].text = re.sub('MARITAL STATUS:','MARITAL STATUS: ' + dc['Marital Status'] ,row.cells[i+1].text)
                        except:
                            pass
                        try:
                            row.cells[i+1].text = re.sub('NATIONALITY:','NATIONALITY: ' + dc['Nationality'] ,row.cells[i+1].text)
                        except:
                            pass                    
                        try:
                            row.cells[i+1].text = re.sub('D.O.B:','D.O.B: ' + dc['Date of Birth'] ,row.cells[i+1].text)
                        except:
                            pass                                            
                except:
                    pass

    # Save the updated document as a new file
    doc.save(save_path)

    print("Conversion Completed...")
