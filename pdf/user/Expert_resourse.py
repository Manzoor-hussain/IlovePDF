import os
import openai
import docx
import docx2txt
import textwrap
import re
from .keys import api_key
print(api_key)
# /Users/manzoorhussain/Documents/ILOVEPDF/pdf_input/Palak_Singh_Formatted_CV.docx

def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def Expert_Resource_Converter(path):
    # print("asli path",path)
    formatted= "/Users/manzoorhussain/Documents/ILOVEPDF/pdf/media/pdf_output/Expert_Resource.docx"
    #un_formatted="/Users/manzoorhussain/Documents/ILOVEPDF/pdf/media/pdf_input/Nicholas_Eager.docx"
    # print("dekho me yha hu",un_formatted)
    doc = docx.Document(path)
    formated_text = docx2txt.process(formatted)
    unformated_text = docx2txt.process(path)
    # print("ye tecxt ha",unformated_text)

    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    
    print("Conversion Started...")
    # llm=OpenAI(temperature=0, max_tokens=1500,openai_api_key=api_key)
    fields_labels = "Name, EDUCATION, CERTIFICATIONS, DOMAIN EXPERTISE, AWARDS, SKILLS, PROFESSIONAL AFFILIATIONS, PUBLICATIONS, VOLUNTEERING, RELEVANT CERTIFICATIONS AND EXPERIENCES, TOOLS, SOFT SKILLS, TECHNICAL SKILLS, SKILLS OVERVIEW, INTERESTS, LANGUAGES, SUMMARY OVERVIEW, EMPLOYMENT SUMMARY , Duration, Client, Role, Company, Technologies Include"
    test_text=f"""

    Do the following tasks on this text: "{unformated_text}":

    1: Extract the information according to these labels {fields_labels}.

    2: Must include all the fields and subfields i.e. fields_labels in the output. 

    3: The output must be in key value format.

    4: Do not include any fields or subfields other than the mentioned.

    5: Do not include email, phone number and personal address.

    6: The template should look like this:


    Name:


    EDUCATION:

    CERTIFICATIONS:

    DOMAIN EXPERTISE:

    AWARDS:

    SKILLS:

    PROFESSIONAL AFFILIATIONS:

    PUBLICATIONS:

    VOLUNTEERING:

    RELEVANT CERTIFICATIONS AND EXPERIENCES :

    TOOLS:

    SOFT SKILLS:

    TECHNICAL SKILLS:

    SKILLS OVERVIEW:

    INTERESTS:

    LANGUAGES:


    SUMMARY OVERVIEW
     Summary overview should be written in the form of paragraph

    EMPLOYMENT SUMMARY 

        Duration: \n Give single line space after Duration and do not include key word Duration
        Client: \n Give single line space after Client
        Role: \n Give single line space after Role
        Company: \n Give single line space after Company
        Technologies Include: \n Give single line space after Technologies Include
        Responsibilities
        add the responsibilities in the form of list and add "*" before every responsibility. For example
        * first responsibility
        * second responsibility
        * third responsibility
        ...
        after responsibilities give double line space

    """

    result = get_completion(test_text)


    txt='\n'+result.replace("\n","\n\n") + '\n'

    name_pattern = r'(Name ?:|EDUCATION ?:|CERTIFICATIONS ?:|DOMAIN EXPERTISE ?:|AWARDS ?:|SKILLS ?:|PROFESSIONAL AFFILIATIONS ?:|PUBLICATIONS ?:|VOLUNTEERING ?:|RELEVANT CERTIFICATIONS AND EXPERIENCES ?:|TOOLS ?:|SOFT SKILLS ?:|TECHNICAL SKILLS ?:|SKILLS OVERVIEW ?:|INTERESTS ?:|LANGUAGES ?:|SUMMARY OVERVIEW ?:|EMPLOYMENT SUMMARY ?:?)'

    try:
        name = re.split(name_pattern,txt)
    except:
        name = ''

    dc = {i:'' for i in ["Name", "EDUCATION", "CERTIFICATIONS", "DOMAIN EXPERTISE", "AWARDS", "SKILLS", "PROFESSIONAL AFFILIATIONS", "PUBLICATIONS", "VOLUNTEERING", "RELEVANT CERTIFICATIONS AND EXPERIENCES", "TOOLS", "SOFT SKILLS", "TECHNICAL SKILLS", "SKILLS OVERVIEW", "INTERESTS","LANGUAGES", "SUMMARY OVERVIEW", "EMPLOYMENT SUMMARY"]}
    for ind,i in enumerate(name):
        try:
            if i.strip(' :') in dc and name[ind+1].strip(' :') not in dc:
                dc[i.strip(' :')] += name[ind+1].strip()
        except:
            pass

    dc["CERTIFICATIONS"] = re.sub('\n? *\n\n *| *\n *','\n', dc["CERTIFICATIONS"]).replace("-","")
    dc["INTERESTS"] = re.sub('\n? *\n\n *| *\n *','\n', dc["INTERESTS"]).replace("-","")
    dc['EMPLOYMENT SUMMARY'] = dc['EMPLOYMENT SUMMARY'].replace("*", "• ")
    dc['SOFT SKILLS'] = dc['SOFT SKILLS'].replace('\n\n','\n')
    dc['RELEVANT CERTIFICATIONS AND EXPERIENCES'] = dc['RELEVANT CERTIFICATIONS AND EXPERIENCES'].replace(" \n\n", "\n")
    dc["SKILLS"] = re.sub('\n? *\n\n *| *\n *','\n',dc["SKILLS"])
    dc['RELEVANT CERTIFICATIONS AND EXPERIENCES'] = dc['RELEVANT CERTIFICATIONS AND EXPERIENCES'].replace(" \n\n", "\n")
    dc['EDUCATION'] = dc['EDUCATION'].replace(":", "\n").replace(",","\n")

    

    # Open the existing document
    doc = docx.Document(formatted)

    # Get the first paragraph
    for i,p in enumerate(doc.paragraphs):
        for key in dc:
            if p.text.strip(' :\n').lower() == key.lower().replace('current ',''):

                if key.lower() in ['education','employment summary']:
    #                 dc[key] = dc[key].replace("\n\n\n\n", "\n\n")
                    doc.paragraphs[i+2].text = re.sub('\n? *\n\n *| *\n *','\n',dc[key])
    #                           
                elif key.lower() in ['certifications','interests', 'languages', 'skills']:
                    formatted_text = ''
                    groups = re.split(',|;|-|\n',str(dc[key]).strip())
    #                 print (groups)
                    for group in groups:
    #                     print(groups)
    #                     break

                        if len(groups) == 1:
                            formatted_text = str(dc[key])
                            break
                        wrapper = textwrap.TextWrapper(width=60, initial_indent='• ',
                                                       subsequent_indent='  ')
                        formatted_text += wrapper.fill(group) + '\n'
                    doc.paragraphs[i+2].text = formatted_text.strip()
    #                 
                elif key.lower() in ['name']:
                    doc.paragraphs[i].text = doc.paragraphs[i].text+" "+str(dc[key])


                else:
                    doc.paragraphs[i+2].text = str(dc[key])

    for table in doc.tables:
        for row in table.rows:
            for i,cell in enumerate(row.cells):
                for key in dc:
                    if cell.text.strip(' :\n').lower() == key.lower().replace('current ',''):
                        row.cells[i+1].text = str(dc[key])

    # Save the updated document as a new file
    doc.save('/Users/manzoorhussain/Documents/ILOVEPDF/pdf/media/output_expert_resource.docx')
    print ("\n")
    print("Conversion Completed...")

#Expert_Resource_Converter()